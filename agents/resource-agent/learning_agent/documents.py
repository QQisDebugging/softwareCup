import csv
import json
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from learning_agent.schemas import KnowledgeDocumentInput, KnowledgeIngestRequest

try:
    from langchain_core.documents import Document as LangChainDocument
except Exception:  # pragma: no cover - exercised when optional package is absent
    LangChainDocument = None  # type: ignore[assignment]

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except Exception:  # pragma: no cover - exercised when optional package is absent
    RecursiveCharacterTextSplitter = None  # type: ignore[assignment]


@dataclass
class KnowledgeDocument:
    id: str
    title: str
    text: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class KnowledgeChunk:
    id: str
    document_id: str
    title: str
    text: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    def __init__(self, project_root: Path) -> None:
        self.project_root = project_root

    def load_seed_documents(self, paths: list[str]) -> list[KnowledgeDocument]:
        documents: list[KnowledgeDocument] = []
        for path in paths:
            resolved = self._resolve_path(path)
            if resolved.exists():
                documents.extend(self.load_path(resolved))
        return documents

    def load_request_documents(self, request: KnowledgeIngestRequest) -> list[KnowledgeDocument]:
        documents: list[KnowledgeDocument] = []
        for path in request.paths:
            documents.extend(self.load_path(self._resolve_path(path)))
        for document in request.documents:
            documents.append(self._from_inline_document(document))
        if not documents:
            raise ValueError("No knowledge documents were provided.")
        return documents

    def load_path(self, path: Path) -> list[KnowledgeDocument]:
        if not path.exists():
            raise FileNotFoundError(f"Knowledge path not found: {path}")
        if path.is_dir():
            documents: list[KnowledgeDocument] = []
            for child in sorted(path.rglob("*")):
                if child.is_file() and self._is_supported(child):
                    documents.extend(self.load_path(child))
            return documents
        if not self._is_supported(path):
            return []
        return [self._parse_file(path)]

    def split_documents(self, documents: list[KnowledgeDocument]) -> list[KnowledgeChunk]:
        if not documents:
            return []
        if RecursiveCharacterTextSplitter is not None and LangChainDocument is not None:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=900,
                chunk_overlap=140,
                separators=["\n\n", "\n", "。", "；", ";", ".", " ", ""],
            )
            langchain_docs = [
                LangChainDocument(
                    page_content=document.text,
                    metadata={
                        "document_id": document.id,
                        "title": document.title,
                        "source": document.source,
                        **document.metadata,
                    },
                )
                for document in documents
            ]
            chunks: list[KnowledgeChunk] = []
            for item in splitter.split_documents(langchain_docs):
                metadata = dict(item.metadata)
                chunk_id = f"{metadata.get('document_id', uuid.uuid4().hex)}:{len(chunks)}"
                chunks.append(
                    KnowledgeChunk(
                        id=chunk_id,
                        document_id=str(metadata.pop("document_id", "")),
                        title=str(metadata.pop("title", "未命名资料")),
                        text=item.page_content.strip(),
                        source=str(metadata.pop("source", "inline")),
                        metadata=metadata,
                    )
                )
            return chunks
        return self._fallback_split(documents)

    def _parse_file(self, path: Path) -> KnowledgeDocument:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".json":
            text = self._json_to_text(json.loads(path.read_text(encoding="utf-8")))
        elif suffix == ".csv":
            text = self._csv_to_text(path)
        elif suffix == ".pdf":
            text = self._pdf_to_text(path)
        elif suffix == ".docx":
            text = self._docx_to_text(path)
        else:
            text = ""
        return KnowledgeDocument(
            id=uuid.uuid5(uuid.NAMESPACE_URL, str(path.resolve())).hex,
            title=path.stem,
            text=text.strip(),
            source=str(path.relative_to(self.project_root)) if path.is_relative_to(self.project_root) else str(path),
            metadata={"fileType": suffix.lstrip(".")},
        )

    def _from_inline_document(self, document: KnowledgeDocumentInput) -> KnowledgeDocument:
        title = document.title or document.id or "inline_document"
        return KnowledgeDocument(
            id=document.id or uuid.uuid4().hex,
            title=title,
            text=document.text.strip(),
            source=document.metadata.get("source", "inline"),
            metadata=document.metadata,
        )

    def _resolve_path(self, path: str) -> Path:
        candidate = Path(path)
        if candidate.is_absolute():
            return candidate
        return (self.project_root / candidate).resolve()

    def _is_supported(self, path: Path) -> bool:
        return path.suffix.lower() in {".txt", ".md", ".markdown", ".json", ".csv", ".pdf", ".docx"}

    def _fallback_split(self, documents: list[KnowledgeDocument]) -> list[KnowledgeChunk]:
        chunks: list[KnowledgeChunk] = []
        for document in documents:
            text = document.text.strip()
            start = 0
            index = 0
            while start < len(text):
                piece = text[start : start + 900].strip()
                if piece:
                    chunks.append(
                        KnowledgeChunk(
                            id=f"{document.id}:{index}",
                            document_id=document.id,
                            title=document.title,
                            text=piece,
                            source=document.source,
                            metadata=document.metadata,
                        )
                    )
                index += 1
                start += 760
        return chunks

    def _json_to_text(self, value: Any, prefix: str = "") -> str:
        if isinstance(value, dict):
            lines: list[str] = []
            for key, item in value.items():
                child_prefix = f"{prefix}{key}"
                if isinstance(item, (dict, list)):
                    lines.append(f"{child_prefix}:")
                    lines.append(self._json_to_text(item, prefix=f"{child_prefix}."))
                else:
                    lines.append(f"{child_prefix}: {item}")
            return "\n".join(line for line in lines if line)
        if isinstance(value, list):
            return "\n".join(self._json_to_text(item, prefix=prefix) for item in value)
        return str(value)

    def _csv_to_text(self, path: Path) -> str:
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as file:
            reader = csv.DictReader(file)
            return "\n".join(json.dumps(row, ensure_ascii=False) for row in reader)

    def _pdf_to_text(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:  # pragma: no cover
            raise ValueError("PDF parsing requires pypdf. Install requirements.txt first.") from exc
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _docx_to_text(self, path: Path) -> str:
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover
            raise ValueError("DOCX parsing requires python-docx. Install requirements.txt first.") from exc
        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

