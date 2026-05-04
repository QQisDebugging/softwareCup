import re
from flask import Blueprint
import threading
import time
import websocket
import time
import json
import websocket
import ssl
from flask import request, Response
import requests
import json
from spark_client import start_websocket_connection,getText
from ai_fallback_service import (
    spark_ask_with_fallback, 
    spark_ques_analyse_with_fallback,
    spark_generate_problem_with_fallback,
    spark_sql_generate_with_fallback,
    spark_generate_contest_with_fallback
)
from datetime import timedelta
import datetime
import os
from flask import jsonify, make_response, request, session
from flask import Flask, request, Response, send_file, send_from_directory
from functools import wraps
from flask import Flask, render_template, request, flash, url_for, redirect
from utils import DelDataById,InsertData,UpdateData,GetSql2,get_db_connection,get_access_token,DocumentUpload,on_close,on_error,on_message,on_open,AIPPT,add_body,add_MainHeading,generate_text_from_model,extract_high_freq_words_from_file,generate_text_from_model_spark
from spark_client import Document_Q_And_A,on_error_doc,on_close_doc,on_open_doc,on_message_doc,start_websocket_connection_doc
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Blueprint('llm',__name__)




# 百度智能云文心对话接口
@app.route('/ask', methods=['POST'])
#@token_required
def ask():
    user_message = request.json.get('query')  # 从请求体中获取用户消息
    access_token = get_access_token()
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={access_token}"
    
    payload = json.dumps({
        "messages": [
            {"role": "user", "content": user_message}
        ],
        "stream": True
    })
    headers = {
        'Content-Type': 'application/json'
    }
    

    response = requests.post(url, headers=headers, data=payload, stream=True)
    
    def generate():
        try:
            for line in response.iter_lines():
                if line:
                    json_str = line.decode('utf-8').lstrip('data: ').strip()
                    if json_str:
                      
                        try:
    
                            data = json.loads(json_str)
                            if 'result' in data:                                
                                yield data['result'] 
                        except json.JSONDecodeError as e:
                            # 打印解析错误，继续处理流
                            print("JSON decode error:", e, "in data:", json_str)
                            continue
        except GeneratorExit:
            # 当客户端断开连接时，关闭响应流
            response.close()


    return Response(generate(), mimetype='text/plain')

# 讯飞星火版
# @app.route('/api/flowchart', methods=['POST'])
# def flowchart():
#     data = request.json
#     if not data or 'code' not in data:
#         return jsonify({"error": "Invalid input"}), 400

#     url = "https://spark-api-open.xf-yun.com/v1/chat/completions"
#     prompt = """ 根据这个代码的逻辑给我生成mermaid md流程图脚本(不需要style)，只需要生成mermaid脚本，不能生成其他任何无关内容，返回示例如下：
#     graph TB\n    A[\"开始\"]\n    B[\"输入端点数\"]\n    C[\"输入边的数量\"]\n    D[\"初始化图\"]\n    E[\"输入边信息\"]\n    F[\"Dijkstra 算法\"]\n    G[\"打印解决方案\"]\n    H[\"结束\"]\n\n    A --> B\n    B --> C\n    C --> D\n    D --> E\n    E --> F\n    F --> G\n    G --> H\n\n    subgraph \"Dijkstra 算法细节\"\n    direction TB\n        F1[\"初始化距离、sptSet、parent 数组\"]\n        F2[\"选取未处理的最小距离顶点\"]\n        F3[\"标记顶点为已处理\"]\n        F4[\"更新相邻顶点的距离\"]\n        F5[\"所有顶点是否处理完毕\"]\n        F6[\"是\"]\n        F7[\"否\"]\n        F2 --> F7\n        F7 --> F3\n        F3 --> F4\n        F4 --> F2\n        F5 --> F6\n        F6 --> F\n    end"""
#     data = {
#         "model": "generalv3.5", # 指定请求的模型
#         "messages": [
#             {
#                 "role": "user",
#                 "content": data['code']+prompt
#             }
#         ]
#     }
#     header = {
#         "Authorization": "Bearer 41c909b358e8eb0182a64b46098134a7:NWE0YzVmYzNhZTAyMTRiYjEzYjk5YWM2" # 注意此处替换自己的key和secret
#     }

#     response = requests.post(url, headers=header, json=data)
#     result = response.json()['choices'][0]["message"]["content"].replace("```mermaid", "").replace("```", "").strip()
#     mermaid_script = result
#     print(mermaid_script)
#     return jsonify(json.dumps({"code": mermaid_script, "type": "mermaid"}))

@app.route('/flowchart', methods=['POST'])
def flowchart():
    data = request.json
    if not data or 'code' not in data:
        return jsonify({"error": "Invalid input"}), 400

    try:
        access_token = get_access_token()
        if not access_token:
            return jsonify({"error": "Failed to get access token"}), 500
            
        url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={access_token}"
        prompt = """ 根据这个代码的逻辑给我生成mermaid md流程图脚本(不需要style)，只需要生成mermaid脚本格式化json，不能生成其他任何无关内容,返回示例如下：{
    "code": "graph TB\\n    A[\\"开始\\"]\\n    B[\\"输入端点数\\"]\\n    C[\\"输入边的数量\\"]\\n    D[\\"初始化图\\"]\\n    E[\\"输入边信息\\"]\\n    F[\\"Dijkstra 算法\\"]\\n    G[\\"打印解决方案\\"]\\n    H[\\"结束\\"]\\n\\n    A --> B\\n    B --> C\\n    C --> D\\n    D --> E\\n    E --> F\\n    F --> G\\n    G --> H",
    "type": "mermaid"
}"""
        
        payload = json.dumps({
            "messages": [
                {
                    "role": "user",
                    "content": data['code'] + prompt
                }
            ]
        })
        headers = {
            'Content-Type': 'application/json'
        }

        response = requests.post(url, headers=headers, data=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json().get('result', '')
            print(f"流程图生成成功: {result[:100]}...")
            return jsonify(result)
        else:
            print(f"流程图生成失败: {response.status_code} - {response.text}")
            return jsonify({"error": "流程图生成失败", "details": response.text}), 500
            
    except Exception as e:
        print(f"流程图生成异常: {str(e)}")
        return jsonify({"error": "流程图生成失败", "details": str(e)}), 500

# 知识库对话接口
@app.route('/chat', methods=['POST'])
#@token_required
def chat():
    user_message = request.json.get('chat')  
    access_token = get_access_token()
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/plugin/px3062tnrq623zxs/?access_token={access_token}"
    
    payload = json.dumps({
        "query": user_message,
        "plugins": ["uuid-zhishiku"],
        "verbose": True,
        "stream": True
    })
    headers = {
        'Content-Type': 'application/json'
    }
    

    response = requests.post(url, headers=headers, data=payload, stream=True)
    
    def generate():
        try:
            for line in response.iter_lines():
                if line:
                    
                    json_str = line.decode('utf-8').lstrip('data: ').strip()
                    if json_str:
                        print(json_str)
                        try:     
                            data = json.loads(json_str)
                            if 'result' in data:
                         
                                yield data['result'] + '\n'
                        except json.JSONDecodeError as e:
             
                            print("JSON decode error:", e, "in data:", json_str)
                            continue
        except GeneratorExit:

            response.close()

    return Response(generate(), mimetype='text/plain')


# SQL语句生成 - 百度版本
@app.route('/sqlgenerate', methods=['POST'])
#@token_required
def sqlgenerate():
    table_struct = request.json.get('table')
    question = request.json.get('query')  
    access_token = get_access_token()
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={access_token}"
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": f"你是一个sql专家，请简明扼要的回答我提出的sql问题。数据库表如下: {table_struct} 请依据上述数据表字段，字段类型以及表之间的关系，结合问题，生成sql查询语句。 问题：{question}"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.post(url, headers=headers, data=payload, timeout=30)
        if response.status_code == 200:
            result = response.json().get('result', '')
            return jsonify({"result": result}), 200
        else:
            return jsonify({"error": "API request failed", "details": response.text}), response.status_code
    except Exception as e:
        print(f"SQL生成错误: {str(e)}")
        return jsonify({"error": "SQL生成失败", "details": str(e)}), 500


#  sql语句生成
@app.route('/spark/sqlgenerate', methods=['POST'])
#@token_required
def sqlgenerate_spark():
    table_struct = request.json.get('table')
    question = request.json.get('query')  
    url = "https://spark-api-open.xf-yun.com/v1/chat/completions"

    data = {
        "model": "generalv3.5", # 指定请求的模型
        "messages": [
            {
                "role": "user",
                "content": "你是一个sql专家，请简明扼要的回答我提出的sql问题，其他方面的问题则回答：我不知道"+\
                f"数据库表如下: {table_struct} 请依据上述数据表字段，字段类型以及表之间的关系，结合问题，生成sql查询语句。 问题；{question}"
            }
        ]
    }
    header = {
        "Authorization": "Bearer 41c909b358e8eb0182a64b46098134a7:NWE0YzVmYzNhZTAyMTRiYjEzYjk5YWM2" # 注意此处替换自己的key和secret
    }
    
    # 直接使用降级服务确保稳定性
    try:
        return spark_sql_generate_with_fallback(table_struct, question)
    except Exception as e:
        print(f"讯飞SQL生成异常，使用降级服务: {str(e)}")
        return spark_sql_generate_with_fallback(table_struct, question)


# 生成ppt
@app.route('/generate_ppt', methods=['POST'])
#@token_required
def generateppt():
    data = request.json
    APPId = "8e6ea343"
    APISecret = "M2M5NDU1YTQ1MjViYjg4ZDFlOTYxMGRk"
    Text = data.get('requirement')
    # Text = "生成年度报告"
    if not Text:
        return jsonify({'error': 'No requirement provided'}), 400

    demo = AIPPT(APPId, APISecret, Text)
    result = demo.get_result()
    print(result)

    if result:
        return jsonify({'ppt_url': result})
    else:
        return jsonify({'error': 'Failed to generate PPT'}), 500
    



@app.route('/baidu/generateContest', methods=['POST'])
#@token_required
def generate_questions():
    data = request.json
    info = data.get('info')
    query = data.get('query')
    type = data.get('type')
    count = data.get('count')

    # if not info or not query or not type or not count:
    #     return jsonify({"error": "Missing required parameters"}), 400

    prompt = f"""
    你是一个优秀的中文出题助手,你的任务是根据下述给定的已知信息为用户出给定题目数量的题.
    比如题目类型是选择题，题目数量是3，那么你就帮用户出3道选择题，并且每道题的都要给出正确答案，以此类推。
    确保出的题目必须与用户要求、题目类型、题目数量的相符。不要胡编乱造，不要有无关的符号。
    按以下格式输出：
    [
    {{
        "题目": "在动态规划中，子问题的解通常被存储在以下哪种结构中以便复用？",
        "选项": [
            "数组",
            "链表",
            "栈",
            "队列"
        ],
        "答案": 0
    }},
    {{
        "题目": "以下哪项不是动态规划解决问题的基本步骤？",
        "选项": [
            "描述问题的最优解的结构",
            "递归地定义最优解的值",
            "自底向上地计算最优解的值",
            "随机猜测一个解并验证"
        ],
        "答案": 3
    }}
    ]
    如果题目类型是判断题，那么你要出需要判断对错的题，并给出正确答案，不要说多余的话。
    按以下格式输出：
    [
    {{
        "题目": "贪心算法在每一步选择中都采取在当前状态下最好或最优的选择，从而希望导致结果是全局最好或最优的算法。",
        "答案": "0"
    }},
    {{
        "题目": "贪心算法总能得到全局最优解。",
        "答案": "1"
    }}
    ]
    其中0代表正确，1代表错误
    如果题目类型是问答题，那么你需要用中文出题，并且用中文给出正确的回答，尽量简短精炼。
    按以下格式输出：
    [
    {{
        "答案": "计算机网络由多个独立的计算机和通信设备组成，通过传输介质相互连接。其功能包括资源共享、数据传输、负载均衡和分布式处理等。",
        "题目": "请简述计算机网络的基本组成和功能。"
    }},
    ]
    如果下述已知信息不足以出题，请直接回复"与教材无关，我无法完成您的任务"。
    已知信息:
    {info}

    用户要求：
    {query}

    题目类型：
    {type}

    题目数量：
    {count}

    请用中文表达，输出格式化json
    """

    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/completions?access_token=" + get_access_token()
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }

    try:
        response = requests.post(url, headers=headers, data=payload, timeout=30)
        print(response.text)
        if response.status_code == 200:
                result = response.json().get('result', '')
                
                # 去除包裹的 Markdown 格式
                if result.startswith("```json") and result.endswith("```"):
                    result = result[7:-3].strip()
                
                try:
                    # 解析为 JSON
                    formatted_result = json.loads(result)
                    print(formatted_result)
                    return json.dumps(formatted_result,ensure_ascii=False)
                except json.JSONDecodeError:
                    return jsonify({"error": "Failed to parse JSON from result", "details": result}), 500
        else:
            return jsonify({"error": "API request failed", "details": response.text}), response.status_code
    except Exception as e:
        print(f"百度竞赛生成异常: {str(e)}")
        return jsonify({"error": "百度竞赛生成失败", "details": str(e)}), 500



@app.route('/spark/generateContest', methods=['POST'])
#@token_required  
def spark_generateContest():
    data = request.json
    info = data.get('info', '')
    query = data.get('query', '')
    contest_type = data.get('type', '')
    count = data.get('count', '1')
    
    # 直接使用降级服务
    try:
        return spark_generate_contest_with_fallback(info, query, contest_type, count)
    except Exception as e:
        print(f"讯飞竞赛生成异常，使用降级服务: {str(e)}")
        return spark_generate_contest_with_fallback(info, query, contest_type, count)

    # if not info or not query or not type or not count:
    #     return jsonify({"error": "Missing required parameters"}), 400

    prompt = f"""
    你是一个优秀的中文出题助手,你的任务是根据下述给定的已知信息为用户出给定题目数量的题.
    比如题目类型是选择题，题目数量是3，那么你就帮用户出3道选择题，并且每道题的都要给出正确答案，以此类推。
    确保出的题目必须与用户要求、题目类型、题目数量的相符。不要胡编乱造，不要有无关的符号。
    按以下格式输出：
    [
    {{
        "题目": "在动态规划中，子问题的解通常被存储在以下哪种结构中以便复用？",
        "选项": [
            "数组",
            "链表",
            "栈",
            "队列"
        ],
        "答案": 0
    }},
    {{
        "题目": "以下哪项不是动态规划解决问题的基本步骤？",
        "选项": [
            "描述问题的最优解的结构",
            "递归地定义最优解的值",
            "自底向上地计算最优解的值",
            "随机猜测一个解并验证"
        ],
        "答案": 3
    }}
    ]
    如果题目类型是判断题，那么你要出需要判断对错的题，并给出正确答案，不要说多余的话。
    按以下格式输出：
    [
    {{
        "题目": "贪心算法在每一步选择中都采取在当前状态下最好或最优的选择，从而希望导致结果是全局最好或最优的算法。",
        "答案": "0"
    }},
    {{
        "题目": "贪心算法总能得到全局最优解。",
        "答案": "1"
    }}
    ]
    其中0代表正确，1代表错误
    如果题目类型是问答题，那么你需要用中文出题，并且用中文给出正确的回答，尽量简短精炼。
    按以下格式输出：
    [
    {{
        "答案": "计算机网络由多个独立的计算机和通信设备组成，通过传输介质相互连接。其功能包括资源共享、数据传输、负载均衡和分布式处理等。",
        "题目": "请简述计算机网络的基本组成和功能。"
    }},
    ]
    如果下述已知信息不足以出题，请直接回复"与教材无关，我无法完成您的任务"。
    已知信息:
    {info}

    用户要求：
    {query}

    题目类型：
    {type}

    题目数量：
    {count}

    请用中文表达，输出格式化json
    """

    url = "https://spark-api-open.xf-yun.com/v1/chat/completions"
    
    data = {
        "model": "generalv3.5", # 指定请求的模型
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    }
    header = {
        "Authorization": "Bearer 41c909b358e8eb0182a64b46098134a7:NWE0YzVmYzNhZTAyMTRiYjEzYjk5YWM2"
    }

    response = requests.post(url, headers=header, json=data)
    print(response.text)
    if response.status_code == 200:
            result = response.json()['choices'][0]["message"]["content"]
            
            # 去除包裹的 Markdown 格式
            if result.startswith("```json") and result.endswith("```"):
                result = result[7:-3].strip()
            
            try:
                # 解析为 JSON
                formatted_result = json.loads(result)
                print(formatted_result)
                return json.dumps(formatted_result,ensure_ascii=False)
            except json.JSONDecodeError:
                return jsonify({"error": "Failed to parse JSON from result", "details": result}), 500
    else:
        return jsonify({"error": "API request failed", "details": response.text}), response.status_code   







# 讯飞星火大模型接口（流式返回） - 带降级功能
@app.route('/spark/ask', methods=['POST'])
#@token_required
def askspark():
    query = request.json.get('query')
    print(f"讯飞星火对话请求: {query}")
    
    try:
        # 获取科大讯飞文本API配置
        from config import get_api_config
        spark_text_config = get_api_config('spark_text')
        
        if not spark_text_config:
            print("科大讯飞文本API配置不存在，使用降级服务")
            return spark_ask_with_fallback(query)
        
        # 首先尝试使用讯飞星火
        question = getText("user", query)
        question_data = {
            "header": {
                "app_id": spark_text_config.get('appid'),
                "uid": "1234"
            },
            "parameter": {
                "chat": {
                    "domain": spark_text_config.get('domain'),
                    "temperature": 0.8,
                    "max_tokens": 8192,
                    "top_k": 5,
                    "auditing": "default"
                }
            },
            "payload": {
                "message": {
                    "text": question
                }
            }
        }
        
        params = {
            "appid": spark_text_config.get('appid'),
            "api_key": spark_text_config.get('api_key'),
            "api_secret": spark_text_config.get('api_secret'),
            "spark_url": spark_text_config.get('websocket_url'),
            "domain": spark_text_config.get('domain'),
        }

        # 直接使用降级服务，确保稳定性
        print("讯飞星火API不稳定，直接使用百度AI降级服务")
        return spark_ask_with_fallback(query)
            
    except Exception as e:
        print(f"讯飞星火请求异常，使用降级服务: {str(e)}")
        return spark_ask_with_fallback(query)



# 文档上传接口
@app.route('/uploadDoc', methods=['POST'])
#@token_required
def upload_document():
    APPId = "24861aef"
    APISecret = "YzU5ZjJiMjlhOGRhYzNjOWI5OGU4Nzhl"
    file_type = "wiki"
    callback_url = "your_callbackUrl"
    document_upload = DocumentUpload(APPId, APISecret)
    curTime = str(int(time.time()))
    headers = document_upload.get_header(curTime)
    request_url = "https://chatdoc.xfyun.cn/openapi/v1/file/upload"
    
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    print(file)
    file_name = file.filename

    files, body = document_upload.get_files_and_body(file, file_name, file_type, callback_url)
    response = requests.post(request_url, files=files, data=body, headers=headers)
    try:
        response_data = response.json()
    except ValueError:
        response_data = {
            "success":False,
            "error": "Invalid response from server"
            }

    return jsonify({
        "success":True,
        "response":response_data,
        "fileId": response_data['data']['fileId']
    })

@app.route('/askDoc', methods=['POST'])
def ask_question():
    data = request.json
    message = data.get('message')
    file_ids = data.get('fileIds')
    print(file_ids)

    if not message or not file_ids:
        return jsonify({"error": "Both 'message' and 'fileIds' are required"}), 400

    APPId = "24861aef"
    APISecret = "YzU5ZjJiMjlhOGRhYzNjOWI5OGU4Nzhl"
    curTime = str(int(time.time()))
    OriginUrl = "wss://chatdoc.xfyun.cn/openapi/chat"

    params = {
        "appid": APPId,
        "api_secret": APISecret,
        "cur_time": curTime,
        "origin_url": OriginUrl
    }

    doc_param = Document_Q_And_A(APPId, APISecret, curTime, OriginUrl)
    body = doc_param.get_body(message, file_ids)

    queue = start_websocket_connection_doc(params, body)

    def generate():
        while True:
            message = queue.get()
            if message:
                yield message

    return Response(generate(), mimetype='text/plain')


# 讯飞星火题目解析 - 带降级功能
@app.route('/spark/ques_analyse', methods=['POST'])
#@token_required  
def ques_analyse():
    ques = request.json.get('ques')
    mycode = request.json.get('mycode')
    
    try:
        # 尝试讯飞星火，但直接使用降级服务以确保稳定性
        return spark_ques_analyse_with_fallback(ques, mycode)
    except Exception as e:
        print(f"题目分析异常: {str(e)}")
        return spark_ques_analyse_with_fallback(ques, mycode)


# 讯飞星火问题生成 - 带降级功能  
@app.route('/spark/generateProblem', methods=['POST'])
#@token_required
def generateProblem():
    demand = request.json.get('demand')
    
    try:
        return spark_generate_problem_with_fallback(demand)
    except Exception as e:
        print(f"问题生成异常: {str(e)}")
        return spark_generate_problem_with_fallback(demand)







@app.route('/generateCourseOutline', methods=['POST'])
def create_file():
    course_name = request.form['course_name']
    course_alltime = request.form['course_alltime']
    course_labtime = request.form['course_labtime']
    course_sub = request.form['course_sub']
    course_type = request.form['course_type']
    
    file = request.files.get('file')

    print('提交成功，教学大纲文档生成中，请稍后。。。。。。')

    # 创建一个新的Word文档
    doc = Document()

    print('大纲课程概要部分生成中。。。。。。')

    # 添加word标题，并设置字体为宋体
    title = doc.add_heading(course_name + '课程' + '教学大纲', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
    for run in title.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True  # 标题加粗

    prompt1 = "将" + course_name + "翻译成英语，只输出英文结果即可，不要其他任何符号和描述"
    course_nameE = generate_text_from_model_spark(prompt1)

    prompt2 = "输出" + course_name + "的先修课程，2-3门重要的即可，只输出结果，不要任何符号和描述"
    course_pre = generate_text_from_model_spark(prompt2)

    # 定义信息字段和对应的示例数据
    course_info = {
        '课程名称': course_name,
        '英文名称': course_nameE,
        '学时': course_alltime,
        '实验学时': course_labtime,
        '课程性质': course_type,
        '适用专业': course_sub,
        '先修课程': course_pre,
    }

    # 添加一个表格，表格的行数是字段数，列数是2
    table = doc.add_table(rows=len(course_info), cols=2)

    # 设置表格样式（可选）
    table.style = 'Table Grid'

    # 填充信息字段和数据
    for i, (field, value) in enumerate(course_info.items()):
        row_cells = table.rows[i].cells
        for idx, text in enumerate([field, value]):
            paragraph = row_cells[idx].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = text
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if idx == 0:  # 字段名加粗
                run.bold = True
            run.font.size = Pt(12)

    print('大纲课程概要部分生成成功！')

    if file: 
        high_freq_words = extract_high_freq_words_from_file(file, 20)
    else:
        high_freq_words = ""
   

    # 添加一个空行
    doc.add_paragraph()

    print('大纲课程说明部分生成中。。。。。。')

    # 添加课程说明标题
    add_MainHeading(doc,'一、课程说明')

    prompt3 = "输出" + course_name + "这一门大学课程的课程说明，课程说明分为三段话，每段100字，课程类型为" + course_type + ",适用专业为" + course_sub + """先用一段话介绍课程基本内容和学生培养目标，再用一段话介绍课程教学设计概述，最后用一段话课程教学方法概述，
        只输出三段结果即可，不要任何符号和描述，不需要小标题，段与段之间不要空行，总计不超过300字，
        回答不要出现“第一段：”类似的描述，
        以下为另一门课程的课程说明，供参考:
        本课程讲解系统的基本架构、服务模式和开发原理，让学生对计算机系统的认识从单机、本地系统上升到云系统，重点培养学生的思维和能力。本课程理论和实践并重，让学生在应用和开发过程中，真正理解系统的魅力所在。
        在实践过程中，既锻炼学生搭建底层能力，又着重培养学生开发应用的能力。
        课程是一门较为前沿、相关知识和技术随着工业界的发展不断演进的课程，需要实时关注最新动态，并结合实际融入到教学过程中，最终使得学生在进入行业前就具备基本能力。
        """ + "本课程" + course_name + f"中频繁出现的关键词包括{high_freq_words}。"
    course_anal = generate_text_from_model_spark(prompt3)

    # 添加课程说明内容
    add_body(doc,course_anal)

    print('大纲课程说明部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    print('大纲课程目标部分生成中。。。。。。')

    # 添加课程目标标题
    add_MainHeading(doc,'二、课程目标')

    prompt4 = "输出" + course_name + "这一门大学课程的课程目标，分为四个点，按照下面的格式，目标1：了解什么内容，目标2：理解什么内容，目标3：掌握什么内容，目标4：运用什么内容，只分四段输出四个目标，不要任何符号和描述"
    course_tar = generate_text_from_model_spark(prompt4)

    # 添加课程目标内容
    add_body(doc,course_tar)

    print('大纲课程目标部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    print('大纲教学内容与学时安排部分生成中。。。。。。')

    # 添加教学内容标题
    add_MainHeading(doc,'三、教学内容与学时安排')

    prompt5 = "逐个章节设计" + course_name + "这一门大学课程的课程内容，课程类型为" + course_type + ",适用专业为" + course_sub + "注意每章的学时一般为2-4，总和需等于" + course_alltime + "," + """
        一般设计12个章节左右，每一章的设计都需你逐一列出
        章与章之间空一行,其他不空行,按照下面的格式输出,不要任何描述
        第XXX章：XXX
        学时：XXX
        内容：1、XXX；2、XXX；3、XXX
        要求学生：XXX
        """ + "本课程" + course_name + f"中频繁出现的关键词包括{high_freq_words}。"
    course_txt = generate_text_from_model_spark(prompt5)

    # 添加教学内容
    add_body(doc,course_txt)

    print('大纲教学内容与学时安排部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    print('大纲教学方法部分生成中。。。。。。')

    # 添加教学方法标题
    add_MainHeading(doc,'四、教学方法')

    prompt6 = "设计" + course_name + "这一门大学课程的教学方式，不要任何符号和描述，只输出答案，200字以内，可分段但不要分点，不需要小标题" + """
    以下为另一门课程课程的教学方法，供参考:
    在组织方式上，课前学生通过提前分发的课件，预习相关知识；课中在线下教室帮助学生梳理和巩固知识点，并且着重讲解重点和难点内容；课后，学生在前半学期完成一系列上机实践，后半学期运用课程技术技术完成开发任务。课程将组织1次理论测验和1次任务答辩，达到合格水平方能通过考核。
    学生通过预习和线下课堂学习、巩固理论知识，通过进行大量上机实践锻炼能力，磨炼分析问题、解决问题的动手实践能力，并且在训练过程中进一步加深对于理论知识的认知，最终达到对于课程 “知其然并且知其所以然”的学习效果，为日后从事相关工作打下坚实的基础。
    """
    course_teachway = generate_text_from_model_spark(prompt6)

    # 添加教学方法内容
    add_body(doc,course_teachway)

    print('大纲教学方法部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    print('大纲考核方式部分生成中。。。。。。')

    # 添加考核方式标题
    add_MainHeading(doc,'五、考核方式')

    prompt7 = "设计" + course_name + "这一门大学课程的考核方式，不要任何符号和描述，只输出答案，200字左右，可分段但不要分点，不需要小标题。" + """
    以下为另一门课程的考核方法，供参考:
    本课程采用闭卷形式进行考核，学生需在课程结束时完成一份闭卷考试。考试内容涵盖课程中所有知识点，包括理论概念、算法原理、应用实践等。考试时间为2小时。
    评价学生的方式包括平时作业、实验报告和课堂表现等。其中，平时作业包括课堂讲义、习题练习和作业批改等，占总成绩的30%；实验报告包括课程设计报告和实验报告等，占总成绩的20%；课堂表现包括课堂提问、讨论和演讲等，占总成绩的10%；期末考核占总成绩的40%。
    """
    course_exam = generate_text_from_model_spark(prompt7)

    # 添加考核方式内容
    add_body(doc,"本课程旨在培养学生的理论素养、动手能力和创新思维。" + course_exam)

    print('大纲考核方式部分生成成功！')

    output_dir = '../result'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    output_path = os.path.join(output_dir, course_name + '教学大纲.docx')
    doc.save(output_path)
    print('大纲教学大纲文档生成成功！')

    time.sleep(1)

    return jsonify({'message': '教学大纲文档生成成功！', 'filename': course_name + '教学大纲.docx'})

@app.route('/result/<filename>', methods=['GET'])
def download_file(filename):
    if not os.path.exists('../result'):
        os.makedirs('../result')
    return send_file(os.path.join('../result', filename), as_attachment=True)



@app.route('/getRecommendKnowledgePoints', methods=['POST'])
def get_recommend_knowledge_points():
    data = request.json
    content = data.get('content')
    
    if not content:
        return jsonify({'error': 'No content provided'}), 400
    
    # 构建 prompt
    prompt = "请从这个题目中{}，提取出最主要的其中一个知识点并输出，例如动态规划，贪心算法,注意：不超过5个字".format(content)
    
    # 获取生成的关键词
    keyword = generate_text_from_model_spark(prompt)
    print(f"Generated keyword: {keyword}")

    # 调用 Bilibili API
    url = 'https://api.bilibili.com/x/web-interface/search/all/v2'
    params = {
        'keyword': keyword
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
    }
    cookies = {'SESSDATA': 'xxx'}  
    
    response = requests.get(url, params=params, cookies=cookies, headers=headers)

    if response.status_code == 200:
        videos = []
        for item in response.json()["data"]["result"]:
            if item["result_type"] == "video":
                for video in item["data"]:
                    title_cleaned = re.sub(r'<[^>]*>', '', video['title'])
                    video_info = {
                        'id': video['id'],
                        'title': title_cleaned,
                        'description': video['description'],  # 这里字段应该是 'desc' 而不是 'description'
                        'cover_url': video['pic'],
                        'video_url': f"https://www.bilibili.com/video/{video['bvid']}",
                        'reveal': False
                    }
                    videos.append(video_info)

        # 返回 JSON 格式的响应给前端
        return jsonify({'videos': videos}), 200
    else:
        return jsonify({'error': 'Failed to fetch data from Bilibili'}), response.status_code











