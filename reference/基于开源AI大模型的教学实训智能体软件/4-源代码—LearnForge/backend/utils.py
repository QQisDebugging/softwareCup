import pymysql
from pymysql import OperationalError
import requests
import hashlib
import base64
import hmac
import time
import random
import os
from urllib.parse import urlencode
import json
import websocket
import _thread as thread
import ssl
import requests
from requests_toolbelt.multipart.encoder import MultipartEncoder
import fitz  # PyMuPDF
from collections import Counter
import jieba
import re

#工具
# 1.大模型diaoyong
# 2.数据库操作
# 3.文档上传



# ppt生成
class AIPPT:
    def __init__(self, APPId, APISecret, Text):
        self.APPid = APPId
        self.APISecret = APISecret
        self.text = Text
        self.header = {}

    def get_signature(self, ts):
        try:
            auth = self.md5(self.APPid + str(ts))
            return self.hmac_sha1_encrypt(auth, self.APISecret)
        except Exception as e:
            print(e)
            return None

    def hmac_sha1_encrypt(self, encrypt_text, encrypt_key):
        return base64.b64encode(hmac.new(encrypt_key.encode('utf-8'), encrypt_text.encode('utf-8'), hashlib.sha1).digest()).decode('utf-8')

    def md5(self, text):
        return hashlib.md5(text.encode('utf-8')).hexdigest()

    def create_task(self):
        url = 'https://zwapi.xfyun.cn/api/aippt/create'
        timestamp = int(time.time())
        signature = self.get_signature(timestamp)
        body = self.get_body(self.text)

        headers = {
            "appId": self.APPid,
            "timestamp": str(timestamp),
            "signature": signature,
            "Content-Type": "application/json; charset=utf-8"
        }
        self.header = headers
        response = requests.post(url, data=json.dumps(body), headers=headers).text
        resp = json.loads(response)
        if resp['code'] == 0:
            return resp['data']['sid']
        else:
            print('创建PPT任务失败')
            return None

    def get_body(self, text):
        body = {
            "query": text
        }
        return body

    def get_process(self, sid):
        print("sid:" + sid)
        if sid:
            response = requests.get(f"https://zwapi.xfyun.cn/api/aippt/progress?sid={sid}", headers=self.header).text
            print(response)
            return response
        else:
            return None

    def get_result(self):
        task_id = self.create_task()
        while True:
            print(1)
            response = self.get_process(task_id)
            resp = json.loads(response)
            process = resp['data']['process']
            if process == 100:
                PPTurl = resp['data']['pptUrl']
                break
            # time.sleep(5)  # 添加一个延迟，避免过于频繁的请求
        return PPTurl
    
#文件上传
class DocumentUpload:
    def __init__(self, APPId, APISecret):
        self.APPId = APPId
        self.APISecret = APISecret

    def get_origin_signature(self, timestamp):
        m2 = hashlib.md5()
        data = bytes(self.APPId + timestamp, encoding="utf-8")
        m2.update(data)
        checkSum = m2.hexdigest()
        return checkSum

    def get_signature(self, timestamp):
        signature_origin = self.get_origin_signature(timestamp)
        signature = hmac.new(self.APISecret.encode('utf-8'), signature_origin.encode('utf-8'),
                             digestmod=hashlib.sha1).digest()
        signature = base64.b64encode(signature).decode(encoding='utf-8')
        return signature

    def get_header(self, timestamp):
        signature = self.get_signature(timestamp)
        header = {
            "appId": self.APPId,
            "timestamp": timestamp,
            "signature": signature,
        }
        return header

    def get_files_and_body(self, file, file_name, file_type, callback_url):
        body = {
            "url": "",
            "fileName": file_name,
            "fileType": file_type,
            "needSummary": False,
            "stepByStep": False,
            "callbackUrl": callback_url,
        }
        files = {'file': (file_name, file.read(), file.content_type)}
        return files, body


def on_error(ws, error):
    print("### error:", error)

def on_close(ws, close_status_code, close_msg):
    print("### closed ###")
    print("关闭代码：", close_status_code)
    print("关闭原因：", close_msg)

def on_open(ws):
    thread.start_new_thread(run, (ws,))

def run(ws, *args):
    data = json.dumps(ws.question)
    ws.send(data)

def on_message(ws, message):
    data = json.loads(message)
    code = data['code']
    if code != 0:
        print(f'请求错误: {code}, {data}')
        ws.close()
    else:
        content = data["content"]
        status = data["status"]
        print(content, end='')
        if status == 2:
            ws.close()


# 大模型调用
def get_access_token():
    """
    使用 API Key，Secret Key 获取 access_token。
    """
    client_id = os.getenv("BAIDU_API_KEY", "")
    client_secret = os.getenv("BAIDU_SECRET_KEY", "")
    if not client_id or not client_secret:
        return None
    url = f"https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id={client_id}&client_secret={client_secret}"
    response = requests.get(url)
    return response.json().get("access_token")


# 数据库配置
def OpenDb():
    # 配置您的 MySQL 数据库连接参数
    config = {
        'host': 'localhost',
        'user': 'a7',
        'password': '123456',
        'database': 'a7',
        'charset': 'utf8mb4',
        'cursorclass': pymysql.cursors.DictCursor  # 使用字典游标读取数据
    }
    conn = pymysql.connect(**config)
    return conn



def get_db_connection():
    try:
        connection = pymysql.connect(
            host='127.0.0.1',
            port=3306,
            user='root',
            password='123456',
            database='a7',
            charset='utf8mb4'
        )
        return connection
    except OperationalError as e:
        print("Error: ", e)
        return None
    
    

def CloseDb(conn):
    conn.close()


def GetSql(conn, sql):
    cur = conn.cursor()
    cur.execute(sql)
    fields = []

    for field in cur.description:
        fields.append(field[0])

    result = cur.fetchall()
    # for item in result:
    #     print(item)
    cur.close()
    return result, fields

def GetSql2(sql):
    conn = OpenDb()
    result, fields = GetSql(conn, sql)
    CloseDb(conn)
    return result, fields


def UpdateData(data, tablename):
    conn = OpenDb()
    values = []
    cursor = conn.cursor()
    idName1 = list(data)[0]
    idName2 = list(data)[1]
    for v in list(data)[2:]:
        values.append("%s='%s'" % (v, data[v]))
    sql = "update %s set %s where %s='%s' and %s='%s'" % (
    tablename, ",".join(values), idName1, data[idName1], idName2, data[idName2])
    # print (sql)
    cursor.execute(sql)
    conn.commit()
    CloseDb(conn)

def InsertData(data, tablename):
    conn = OpenDb()
    cursor = conn.cursor()
    fields = ', '.join(data.keys())
    placeholders = ', '.join(['%s'] * len(data))
    sql = f"INSERT INTO {tablename} ({fields}) VALUES ({placeholders})"
    cursor.execute(sql, list(data.values()))
    conn.commit()
    CloseDb(conn)

def DelDataById(id1, id2, value1, value2, tablename):
    conn = OpenDb()
    cursor = conn.cursor()
    sql = f"DELETE FROM {tablename} WHERE {id1}=%s AND {id2}=%s"
    cursor.execute(sql, (value1, value2))
    conn.commit()
    CloseDb(conn)



# erine_speed基于prompt生成内容
def generate_text_from_model(prompt): 
    access_token = get_access_token()
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={access_token}"
    
   
    payload = json.dumps({
        "system": "You are a helpful assistant",
        "max_output_tokens": 2048,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    response = requests.request("POST", url, headers=headers, data=payload)
    response_json = response.json()
    
    # Assuming the 'result' field exists in the JSON response
    result = response_json.get("result", "")
    # print(result)
    return result

def generate_text_from_model_spark(prompt): 
    try:
        url = "https://spark-api-open.xf-yun.com/v1/chat/completions"
        
        data = {
            "model": "generalv3.5", 
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        }
        header = {
            "Authorization": "Bearer 41c909b358e8eb0182a64b46098134a7:NWE0YzVmYzNhZTAyMTRiYjEzYjk5YWM2" 
        }

        # 添加超时设置
        response = requests.post(url, headers=header, json=data, timeout=30)
        response.raise_for_status()  # 检查HTTP状态码
        
        response_json = response.json()
        print("AI API响应:", response_json)  # 调试日志
        
        # 检查响应格式
        if 'choices' not in response_json or not response_json['choices']:
            print("AI API响应格式错误: 缺少choices字段")
            raise ValueError("AI API响应格式错误")
        
        result = response_json['choices'][0]["message"]["content"]
        print("AI生成结果:", result)  # 调试日志
        return result
        
    except requests.exceptions.Timeout:
        print("AI API请求超时")
        # 返回示例数据作为fallback
        return '''[
    {
        "知识点": "基本输入输出",
        "题目": "编写一个C语言程序，要求用户输入两个整数，然后输出它们的和。使用scanf和printf函数完成输入输出操作。"
    },
    {
        "知识点": "变量和数据类型", 
        "题目": "声明int、float、char类型的变量，为它们赋值，然后使用printf函数按指定格式输出这些变量的值。"
    },
    {
        "知识点": "条件语句",
        "题目": "编写程序判断用户输入的数字是正数、负数还是零，使用if-else语句实现条件判断。"
    }
]'''
    except requests.exceptions.RequestException as e:
        print(f"AI API请求失败: {e}")
        # 返回示例数据作为fallback
        return '''[
    {
        "知识点": "基本输入输出",
        "题目": "编写一个C语言程序，要求用户输入两个整数，然后输出它们的和。使用scanf和printf函数完成输入输出操作。"
    },
    {
        "知识点": "变量和数据类型", 
        "题目": "声明int、float、char类型的变量，为它们赋值，然后使用printf函数按指定格式输出这些变量的值。"
    },
    {
        "知识点": "循环语句",
        "题目": "使用for循环编写程序，计算1到100之间所有整数的和，并输出结果。"
    }
]'''
    except Exception as e:
        print(f"AI服务发生未知错误: {e}")
        # 返回示例数据作为fallback
        return '''[
    {
        "知识点": "基本输入输出",
        "题目": "编写一个C语言程序，要求用户输入两个整数，然后输出它们的和。"
    },
    {
        "知识点": "变量和数据类型", 
        "题目": "声明不同类型的变量并为它们赋值，然后打印输出。"
    }
]'''


# 提取pdf关键词
def extract_high_freq_words_from_file(file_obj, top_n=20):
 
    stop_words = set(["的", "是", "在", "我", "有", "和", "就", "等", "可以", "了", "进行", "对", "都","功能","不同","图","认","人","从","型","时","结构","例如","数","描述","表示","于","类型","要",
                      "为", "与", "中", "通过", "使用", "上", "一个", "共", "或", "并", "将", "也","包括","年","指","被","不","一种","地","大","后","需要","应用","根据","方式","如图","请","步骤","一",
                        "它", "人们", "主要", "如", "单元","章节", "第", "由","下列","方面","一旦","而","以","其","具有","可","到","能","各种","所示","方法""常用","来","例","用","会","个","问题","分析"])

    def clean_and_split_text(text):
        words = jieba.cut(text)
        filtered_words = [word for word in words if word not in stop_words and re.match(r'[\u4e00-\u9fa5]+', word)]
        return filtered_words

    # # 读取文件对象到内存
    file_content = file_obj.read()

    try:
        # 从内存中加载PDF
        doc = fitz.open(stream=file_content, filetype="pdf")
    except Exception as e:
        print("无法从文件对象加载PDF")
        print(e)
        return ""

    full_text = ""
    for page in doc:
        full_text += page.get_text()

    words = clean_and_split_text(full_text)
    word_counts = Counter(words)
    most_common_words = word_counts.most_common(top_n)

    doc.close()

    # 仅返回高频词，用顿号隔开
    return '、'.join([word for word, _ in most_common_words])

from docx.shared import Pt
from docx.oxml.ns import qn
# 添加word一级标题
def add_MainHeading(doc,text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 添加word正文内容
def add_body(doc,text):
    content = text
    p = doc.add_paragraph(content)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = '宋体'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
